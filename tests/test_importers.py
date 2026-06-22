"""Tests for Word and Excel importers."""

import pandas as pd
from docx import Document

from core import importers


def _make_docx(path):
    doc = Document()
    doc.add_paragraph("Lentiviral Transduction Protocol")
    doc.add_paragraph("Thaw cells", style="List Number")
    doc.add_paragraph("Add polybrene", style="List Number")
    doc.add_paragraph("3. Spinoculate at 1000g")
    doc.save(path)


def test_parse_word(tmp_path):
    p = tmp_path / "protocol.docx"
    _make_docx(str(p))
    parsed = importers.parse_word(str(p))
    assert parsed.title == "Lentiviral Transduction Protocol"
    assert "Thaw cells" in parsed.steps
    assert "Spinoculate at 1000g" in parsed.steps  # numeric prefix stripped
    assert parsed.source_filename == "protocol.docx"


def test_import_protocol_persists_steps(conn, tmp_path):
    p = tmp_path / "protocol.docx"
    _make_docx(str(p))
    parsed = importers.parse_word(str(p))
    pid = importers.import_protocol(conn, parsed, tags="viral")
    steps = conn.execute(
        "SELECT count(*) FROM protocol_steps WHERE protocol_id=?", (pid,)
    ).fetchone()[0]
    assert steps == 3


def test_parse_word_includes_table_content(tmp_path):
    """Body text and steps inside tables must not be dropped (the partial-doc bug)."""
    p = tmp_path / "tabled.docx"
    doc = Document()
    doc.add_paragraph("Tabled Protocol")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Step"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "1. Add buffer to tube"
    table.cell(1, 1).text = "10 mL"
    doc.save(str(p))

    parsed = importers.parse_word(str(p))
    assert "Add buffer to tube" in parsed.body_text
    assert "10 mL" in parsed.body_text
    assert "Add buffer to tube" in parsed.steps  # numbered prefix stripped


def test_parse_and_import_excel(conn, tmp_path):
    p = tmp_path / "plan.xlsx"
    pd.DataFrame(
        {
            "Task": ["Transfect", "Harvest"],
            "Date": ["2026-06-16", "2026-06-18"],
            "Sample": ["A1", "A1"],
            "Reagent": ["PEI", None],
        }
    ).to_excel(p, index=False)

    parsed = importers.parse_excel(str(p))
    assert parsed["columns"] == ["Task", "Date", "Sample", "Reagent"]
    assert len(parsed["rows"]) == 2

    mapping = {
        "task_name": "Task",
        "planned_date": "Date",
        "sample": "Sample",
        "reagent": "Reagent",
        "notes": None,
    }
    xid = importers.import_experiment(
        conn, "Virus prep", parsed["rows"], mapping, source_filename="plan.xlsx"
    )
    tasks = conn.execute(
        "SELECT * FROM experiment_tasks WHERE experiment_id=? ORDER BY id", (xid,)
    ).fetchall()
    assert len(tasks) == 2
    assert tasks[0]["task_name"] == "Transfect"
    assert tasks[0]["reagent"] == "PEI"
    assert tasks[1]["reagent"] is None
