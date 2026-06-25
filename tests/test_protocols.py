"""Tests for the protocol viewer: file storage, file serving, and rename."""

import pytest
from docx import Document
from fastapi.testclient import TestClient

from core import files


@pytest.fixture()
def client(tmp_path, monkeypatch):
    """A TestClient whose config points at a throwaway data dir (unlocked)."""
    import server.config as cfg

    monkeypatch.setattr(cfg, "CONFIG_PATH", str(tmp_path / "cfg.json"))
    monkeypatch.setattr(cfg, "DEFAULT_DB_PATH", str(tmp_path / "seshat.db"))
    monkeypatch.setattr(cfg, "DEFAULT_INBOX_DIR", str(tmp_path / "inbox"))

    from server.main import app

    c = TestClient(app)
    c.post("/api/auth/unlock", json={"passphrase": "test-pass", "confirm": "test-pass"})
    return c


def _make_docx(path):
    doc = Document()
    doc.add_paragraph("Lentiviral Transduction Protocol")
    doc.add_paragraph("Thaw cells", style="List Number")
    doc.add_paragraph("Add polybrene", style="List Number")
    doc.save(path)


def test_import_stores_file_bytes(conn, tmp_path):
    p = tmp_path / "protocol.docx"
    _make_docx(str(p))
    pid = files.import_protocol_file(conn, str(p))
    row = conn.execute(
        "SELECT file_data, file_mime FROM protocols WHERE id=?", (pid,)
    ).fetchone()
    assert row["file_data"] is not None
    assert len(bytes(row["file_data"])) > 0
    assert "wordprocessingml" in row["file_mime"]


def test_pdf_mime_detected(conn, tmp_path):
    # Build a real (if blank) PDF with pypdf's writer so parse_pdf succeeds.
    from pypdf import PdfWriter

    p = tmp_path / "proto.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with open(p, "wb") as fh:
        writer.write(fh)

    pid = files.import_protocol_file(conn, str(p))
    row = conn.execute(
        "SELECT file_mime, file_data FROM protocols WHERE id=?", (pid,)
    ).fetchone()
    assert row["file_mime"] == "application/pdf"
    assert bytes(row["file_data"]).startswith(b"%PDF")


def test_rename_and_list_via_api(client, tmp_path):
    # Stage + commit a docx through the upload flow so it has stored bytes.
    p = tmp_path / "myproto.docx"
    _make_docx(str(p))
    with open(p, "rb") as fh:
        up = client.post(
            "/api/files/upload",
            files={"file": ("myproto.docx", fh, "application/octet-stream")},
        )
    staged = up.json()["path"]
    commit = client.post(
        "/api/files/protocol/commit",
        json={"path": staged, "title": "Original Title"},
    )
    pid = commit.json()["id"]

    # List reports has_file=True and excludes the BLOB payload.
    protos = client.get("/api/protocols").json()
    target = next(p for p in protos if p["id"] == pid)
    assert target["has_file"] is True
    assert "file_data" not in target

    # Rename works and is reflected in the list.
    assert client.put(f"/api/protocols/{pid}", json={"title": "Renamed Protocol"}).status_code == 200
    protos = client.get("/api/protocols").json()
    assert next(p for p in protos if p["id"] == pid)["title"] == "Renamed Protocol"

    # Empty title rejected.
    assert client.put(f"/api/protocols/{pid}", json={"title": "  "}).status_code == 400


def test_serve_docx_raw_bytes(client, tmp_path):
    """The viewer endpoint serves the original .docx bytes; the browser (docx-preview)
    renders it client-side, so no server-side HTML conversion happens anymore."""
    p = tmp_path / "viewme.docx"
    _make_docx(str(p))
    with open(p, "rb") as fh:
        up = client.post(
            "/api/files/upload",
            files={"file": ("viewme.docx", fh, "application/octet-stream")},
        )
    staged = up.json()["path"]
    pid = client.post(
        "/api/files/protocol/commit", json={"path": staged, "title": "Viewer"}
    ).json()["id"]

    r = client.get(f"/api/protocols/{pid}/file")
    assert r.status_code == 200
    assert "wordprocessingml" in r.headers["content-type"]
    assert r.content.startswith(b"PK")  # .docx is a zip archive


def _commit_docx(client, tmp_path, name="pdf.docx", title="PDF Me"):
    p = tmp_path / name
    _make_docx(str(p))
    with open(p, "rb") as fh:
        up = client.post(
            "/api/files/upload",
            files={"file": (name, fh, "application/octet-stream")},
        )
    return client.post(
        "/api/files/protocol/commit", json={"path": up.json()["path"], "title": title}
    ).json()["id"]


def test_pdf_endpoint_converts_and_caches(client, tmp_path, monkeypatch):
    """DOCX→PDF is converted on first view and cached; the second view reuses it."""
    from core import docx_pdf

    pid = _commit_docx(client, tmp_path)

    calls = {"n": 0}

    def fake_convert(data):
        calls["n"] += 1
        return b"%PDF-1.4 fake-pdf-bytes"

    monkeypatch.setattr(docx_pdf, "converter_available", lambda: True)
    monkeypatch.setattr(docx_pdf, "convert_docx_to_pdf", fake_convert)

    r1 = client.get(f"/api/protocols/{pid}/file.pdf")
    assert r1.status_code == 200
    assert r1.headers["content-type"] == "application/pdf"
    assert r1.content == b"%PDF-1.4 fake-pdf-bytes"
    assert calls["n"] == 1

    # Second request is served from the cached pdf_render — no re-conversion.
    r2 = client.get(f"/api/protocols/{pid}/file.pdf")
    assert r2.status_code == 200
    assert r2.content == b"%PDF-1.4 fake-pdf-bytes"
    assert calls["n"] == 1


def test_pdf_endpoint_503_without_converter(client, tmp_path, monkeypatch):
    """When no converter is installed, the endpoint signals 503 so the UI falls back."""
    from core import docx_pdf

    pid = _commit_docx(client, tmp_path, name="noconv.docx")
    monkeypatch.setattr(docx_pdf, "converter_available", lambda: False)

    r = client.get(f"/api/protocols/{pid}/file.pdf")
    assert r.status_code == 503


def test_pdf_endpoint_passthrough_for_pdf_protocol(client, tmp_path):
    """A protocol whose original file is already a PDF is served straight through."""
    from pypdf import PdfWriter

    p = tmp_path / "native.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with open(p, "wb") as fh:
        writer.write(fh)
    with open(p, "rb") as fh:
        up = client.post(
            "/api/files/upload", files={"file": ("native.pdf", fh, "application/pdf")}
        )
    pid = client.post(
        "/api/files/protocol/commit", json={"path": up.json()["path"], "title": "Native PDF"}
    ).json()["id"]

    r = client.get(f"/api/protocols/{pid}/file.pdf")
    assert r.status_code == 200
    assert r.headers["content-type"] == "application/pdf"
    assert r.content.startswith(b"%PDF")


def test_update_steps_replaces_all(client, tmp_path):
    p = tmp_path / "edit.docx"
    _make_docx(str(p))
    with open(p, "rb") as fh:
        up = client.post(
            "/api/files/upload",
            files={"file": ("edit.docx", fh, "application/octet-stream")},
        )
    pid = client.post(
        "/api/files/protocol/commit", json={"path": up.json()["path"], "title": "Editable"}
    ).json()["id"]

    # Replace with an edited/reordered/added set; blanks are dropped.
    r = client.put(
        f"/api/protocols/{pid}/steps",
        json={"steps": ["First step", "  ", "Second step", "Third step"]},
    )
    assert r.status_code == 200
    assert r.json()["count"] == 3

    target = next(x for x in client.get("/api/protocols").json() if x["id"] == pid)
    assert [s["text"] for s in target["steps"]] == ["First step", "Second step", "Third step"]
    assert [s["step_no"] for s in target["steps"]] == [1, 2, 3]

    # Unknown protocol → 404.
    assert client.put("/api/protocols/999999/steps", json={"steps": ["x"]}).status_code == 404


def test_file_404_when_no_bytes(conn, client):
    # A protocol imported without file bytes (legacy) returns 404 on /file.
    from core import importers
    from core.models import ParsedProtocol

    # Reach into the live session's connection to insert a fileless protocol.
    import server.security as security

    sess = next(iter(security._sessions.values()))
    pid = importers.import_protocol(
        sess.conn, ParsedProtocol(title="No File", body_text="x", steps=[])
    )
    assert client.get(f"/api/protocols/{pid}/file").status_code == 404
