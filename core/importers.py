"""Importers for existing Word protocols and Excel experiment plans.

Parsing functions are pure (no DB) so they are easy to unit-test. Separate
``import_*`` functions persist the parsed data. Excel layout is unknown ahead
of time, so :func:`parse_excel` simply surfaces the columns and rows and the
UI lets the user map them onto experiment-task fields.
"""

from __future__ import annotations

import json
import os
import re
from typing import Any, Optional

from .db import Connection
from .models import ParsedProtocol
from .notebook import now_ts

# ── Shared text → protocol extraction ───────────────────────────────────────
_NUMBERED = re.compile(r"^\s*(\d+[.)]|[-*•])\s+")


def extract_protocol_from_text(text: str, filename: str) -> ParsedProtocol:
    """Build a ParsedProtocol from raw text (used for PDF and txt/md).

    Title = first non-empty line; steps = lines that look numbered/bulleted.
    """
    lines = [ln.strip() for ln in (text or "").splitlines()]
    nonempty = [ln for ln in lines if ln]
    title = nonempty[0] if nonempty else os.path.splitext(os.path.basename(filename))[0]
    steps = [_NUMBERED.sub("", ln).strip() for ln in nonempty if _NUMBERED.match(ln)]
    return ParsedProtocol(
        title=title,
        body_text="\n".join(nonempty),
        steps=steps,
        source_filename=os.path.basename(filename),
    )


def parse_pdf(path: str) -> ParsedProtocol:
    """Extract text from a PDF protocol (local, offline) via pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(path)
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    return extract_protocol_from_text(text, path)


def parse_text(path: str) -> ParsedProtocol:
    """Parse a plain-text or Markdown protocol."""
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return extract_protocol_from_text(fh.read(), path)


# ── Word (.docx) ────────────────────────────────────────────────────────────
def parse_word(path: str) -> ParsedProtocol:
    """Extract title, full body text and step list from a .docx protocol."""
    from docx import Document  # imported lazily so core import stays light

    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    title = paras[0] if paras else os.path.splitext(os.path.basename(path))[0]

    steps: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if "list" in style or _NUMBERED.match(text):
            steps.append(_NUMBERED.sub("", text).strip())

    return ParsedProtocol(
        title=title,
        body_text="\n".join(paras),
        steps=steps,
        source_filename=os.path.basename(path),
    )


def import_protocol(
    conn: Connection,
    parsed: ParsedProtocol,
    *,
    tags: Optional[str] = None,
) -> int:
    """Persist a parsed protocol and its steps; returns the protocol id."""
    cur = conn.execute(
        """
        INSERT INTO protocols (title, source_filename, version, imported_at, body_text, tags,
                               file_data, file_mime)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            parsed.title,
            parsed.source_filename,
            parsed.version,
            now_ts(),
            parsed.body_text,
            tags,
            parsed.file_data,
            parsed.file_mime,
        ),
    )
    protocol_id = int(cur.lastrowid)
    for i, step in enumerate(parsed.steps, start=1):
        conn.execute(
            "INSERT INTO protocol_steps (protocol_id, step_no, text) VALUES (?, ?, ?)",
            (protocol_id, i, step),
        )
    conn.commit()
    return protocol_id


# ── Excel (.xlsx) ───────────────────────────────────────────────────────────
def parse_excel(path: str, sheet_name: Optional[str] = None) -> dict[str, Any]:
    """Read a sheet and return its columns + rows for column-mapping in the UI.

    Returns ``{"sheet": str, "sheet_names": [...], "columns": [...],
    "rows": [ {col: value, ...}, ... ]}``.
    """
    import math

    import pandas as pd

    xls = pd.ExcelFile(path, engine="openpyxl")
    sheet = sheet_name or xls.sheet_names[0]
    df = xls.parse(sheet)

    # NaN -> None for clean JSON/SQL. (df.where(..., None) is unreliable because
    # pandas treats other=None as "use the default NaN", so clean the records.)
    rows = df.to_dict(orient="records")
    for row in rows:
        for key, val in row.items():
            if val is None or (isinstance(val, float) and math.isnan(val)):
                row[key] = None

    return {
        "sheet": sheet,
        "sheet_names": list(xls.sheet_names),
        "columns": [str(c) for c in df.columns],
        "rows": rows,
    }


# Canonical experiment-task fields the UI maps spreadsheet columns onto.
TASK_FIELDS = ["task_name", "planned_date", "sample", "reagent", "notes"]


def import_experiment(
    conn: Connection,
    name: str,
    rows: list[dict],
    mapping: dict[str, Optional[str]],
    *,
    description: Optional[str] = None,
    planned_date: Optional[str] = None,
    source_filename: Optional[str] = None,
) -> int:
    """Create an experiment and its tasks from mapped spreadsheet rows.

    ``mapping`` maps each canonical field in :data:`TASK_FIELDS` to a source
    column name (or ``None`` to leave it blank).
    """
    cur = conn.execute(
        """
        INSERT INTO experiments
            (name, description, planned_date, status, source_filename, imported_at, metadata_json)
        VALUES (?, ?, ?, 'planned', ?, ?, ?)
        """,
        (
            name,
            description,
            planned_date,
            source_filename,
            now_ts(),
            json.dumps({"mapping": mapping, "row_count": len(rows)}),
        ),
    )
    experiment_id = int(cur.lastrowid)

    def pick(row: dict, field: str) -> Optional[str]:
        col = mapping.get(field)
        if not col or col not in row:
            return None
        val = row[col]
        return None if val is None else str(val)

    for row in rows:
        task_name = pick(row, "task_name") or "(unnamed task)"
        conn.execute(
            """
            INSERT INTO experiment_tasks
                (experiment_id, task_name, planned_date, sample, reagent, notes, status)
            VALUES (?, ?, ?, ?, ?, ?, 'pending')
            """,
            (
                experiment_id,
                task_name,
                pick(row, "planned_date"),
                pick(row, "sample"),
                pick(row, "reagent"),
                pick(row, "notes"),
            ),
        )
    conn.commit()
    return experiment_id
