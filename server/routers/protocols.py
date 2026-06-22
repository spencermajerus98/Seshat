"""Protocols — list/search, rename, delete, and original-file viewer."""

from __future__ import annotations

import html as _html
import io
import re
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import HTMLResponse, Response
from pydantic import BaseModel

from .. import security
from ..deps import current_session, db as db_lock

router = APIRouter(prefix="/api/protocols", tags=["protocols"])

# Columns that are safe to return in list responses (excludes the BLOB).
_LIST_COLS = (
    "id, title, source_filename, version, imported_at, body_text, tags, file_mime, "
    "(file_data IS NOT NULL) AS has_file"
)

_NUMBERED = re.compile(r"^\s*\d+[.)]\s+")
_BULLETED = re.compile(r"^\s*[-*•]\s+")


def _docx_to_html(data: bytes) -> str:
    """Convert a DOCX file (as bytes) to a simple HTML document for in-browser display.

    Walks the document in reading order, rendering both paragraphs and tables.
    ``doc.paragraphs`` alone skips table content — the cause of protocols that
    rendered only partially in the viewer.
    """
    from docx import Document  # lazy import
    from docx.table import Table

    doc = Document(io.BytesIO(data))
    parts = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'>",
        "<style>body{font-family:sans-serif;max-width:860px;margin:0 auto;"
        "padding:1.5rem;line-height:1.6;color:#333}"
        "h1,h2,h3{margin-top:1.2em}p{margin:.4em 0}"
        "ul,ol{margin:.4em 0;padding-left:1.5em}"
        "table{border-collapse:collapse;width:100%;margin:1em 0}"
        "td,th{border:1px solid #ccc;padding:.4em .6em;vertical-align:top;text-align:left}"
        "</style>",
        "</head><body>",
    ]

    # Mutable list state shared across paragraph rendering.
    state = {"in_list": False, "list_tag": "ul"}

    def close_list() -> None:
        if state["in_list"]:
            parts.append(f"</{state['list_tag']}>")
            state["in_list"] = False

    def render_paragraph(p) -> None:
        text = p.text.strip()
        if not text:
            close_list()
            return
        style_name = (p.style.name or "").lower() if p.style else ""
        is_numbered = bool(_NUMBERED.match(text))
        is_bullet = bool(_BULLETED.match(text))
        is_list = is_numbered or is_bullet or "list" in style_name
        e = _html.escape(text)

        if "heading 1" in style_name:
            close_list()
            parts.append(f"<h1>{e}</h1>")
        elif "heading 2" in style_name:
            close_list()
            parts.append(f"<h2>{e}</h2>")
        elif "heading 3" in style_name or "heading 4" in style_name:
            close_list()
            parts.append(f"<h3>{e}</h3>")
        elif is_list:
            new_tag = "ol" if is_numbered else "ul"
            if not state["in_list"]:
                state["list_tag"] = new_tag
                parts.append(f"<{new_tag}>")
                state["in_list"] = True
            elif new_tag != state["list_tag"]:
                parts.append(f"</{state['list_tag']}>")
                state["list_tag"] = new_tag
                parts.append(f"<{new_tag}>")
            item = _NUMBERED.sub("", _BULLETED.sub("", text)).strip()
            parts.append(f"<li>{_html.escape(item)}</li>")
        else:
            close_list()
            parts.append(f"<p>{e}</p>")

    def render_table(tbl) -> None:
        close_list()
        parts.append("<table>")
        for row in tbl.rows:
            parts.append("<tr>")
            for cell in row.cells:
                cell_html = "<br>".join(
                    _html.escape(cp.text.strip())
                    for cp in cell.paragraphs
                    if cp.text and cp.text.strip()
                )
                parts.append(f"<td>{cell_html}</td>")
            parts.append("</tr>")
        parts.append("</table>")

    blocks = (
        doc.iter_inner_content()
        if hasattr(doc, "iter_inner_content")
        else doc.paragraphs
    )
    for block in blocks:
        if isinstance(block, Table):
            render_table(block)
        else:
            render_paragraph(block)

    close_list()
    parts.append("</body></html>")
    return "\n".join(parts)


@router.get("")
def list_protocols(
    q: Optional[str] = None, sess: security.Session = Depends(current_session)
):
    with db_lock(sess) as conn:
        if q:
            like = f"%{q}%"
            rows = conn.execute(
                f"SELECT {_LIST_COLS} FROM protocols "
                "WHERE title LIKE ? OR body_text LIKE ? "
                "ORDER BY title COLLATE NOCASE",
                (like, like),
            ).fetchall()
        else:
            rows = conn.execute(
                f"SELECT {_LIST_COLS} FROM protocols ORDER BY title COLLATE NOCASE"
            ).fetchall()

        protocols = []
        for p in rows:
            steps = conn.execute(
                "SELECT step_no, text FROM protocol_steps WHERE protocol_id=? "
                "ORDER BY step_no",
                (p["id"],),
            ).fetchall()
            d = dict(p)
            d["has_file"] = bool(d.get("has_file", False))
            d["steps"] = [dict(s) for s in steps]
            protocols.append(d)
    return protocols


class ProtocolRename(BaseModel):
    title: str


@router.put("/{protocol_id}")
def rename_protocol(
    protocol_id: int,
    body: ProtocolRename,
    sess: security.Session = Depends(current_session),
):
    if not body.title.strip():
        raise HTTPException(400, "Title is required.")
    with db_lock(sess) as conn:
        conn.execute(
            "UPDATE protocols SET title=? WHERE id=?",
            (body.title.strip(), protocol_id),
        )
        conn.commit()
    return {"ok": True}


class StepsUpdate(BaseModel):
    steps: list[str]


@router.put("/{protocol_id}/steps")
def update_steps(
    protocol_id: int,
    body: StepsUpdate,
    sess: security.Session = Depends(current_session),
):
    """Replace a protocol's steps with the supplied list (edit/add/delete/reorder).

    Blank entries are dropped and the remaining steps are renumbered 1..N.
    """
    cleaned = [s.strip() for s in body.steps if s and s.strip()]
    with db_lock(sess) as conn:
        exists = conn.execute(
            "SELECT 1 FROM protocols WHERE id=?", (protocol_id,)
        ).fetchone()
        if not exists:
            raise HTTPException(404, "Protocol not found.")
        conn.execute("DELETE FROM protocol_steps WHERE protocol_id=?", (protocol_id,))
        for i, text in enumerate(cleaned, start=1):
            conn.execute(
                "INSERT INTO protocol_steps (protocol_id, step_no, text) VALUES (?, ?, ?)",
                (protocol_id, i, text),
            )
        conn.commit()
    return {"ok": True, "count": len(cleaned)}


@router.get("/{protocol_id}/file")
def get_protocol_file(
    protocol_id: int, sess: security.Session = Depends(current_session)
):
    """Return the original file stored at import time.

    PDFs are returned as-is (browser renders them inline in an iframe).
    DOCX files are converted to HTML on the fly.
    Plain-text files are returned as text/plain.
    """
    with db_lock(sess) as conn:
        row = conn.execute(
            "SELECT file_data, file_mime, title FROM protocols WHERE id=?",
            (protocol_id,),
        ).fetchone()

    if not row or not row["file_data"]:
        raise HTTPException(404, "No original file stored for this protocol.")

    data: bytes = bytes(row["file_data"])
    mime: str = row["file_mime"] or "application/octet-stream"

    if "wordprocessingml" in mime:
        return HTMLResponse(content=_docx_to_html(data))

    return Response(content=data, media_type=mime)


@router.delete("/{protocol_id}")
def delete_protocol(protocol_id: int, sess: security.Session = Depends(current_session)):
    with db_lock(sess) as conn:
        conn.execute("DELETE FROM protocols WHERE id=?", (protocol_id,))
        conn.commit()
    return {"ok": True}
